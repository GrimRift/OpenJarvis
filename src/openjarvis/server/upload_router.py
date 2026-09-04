"""Upload / Paste router for ingesting documents into the knowledge store."""

from __future__ import annotations

import asyncio
import io
import logging
import uuid
from typing import List, Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel

from openjarvis.connectors.store import KnowledgeStore
from openjarvis.core.config import DEFAULT_CONFIG_DIR
from openjarvis.server import document_read

logger = logging.getLogger(__name__)

#: Largest attachment accepted. Generous for a paper or a report, bounded
#: so one file cannot occupy the whole request.
MAX_ATTACHMENT_BYTES = 25 * 1024 * 1024

router = APIRouter(prefix="/v1/connectors/upload", tags=["upload"])

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ALLOWED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx"}


def _chunk_text(text: str, max_chars: int = 1000) -> List[str]:
    """Split *text* into ~max_chars pieces at paragraph boundaries."""
    paragraphs = text.split("\n\n")
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if current and len(current) + len(para) + 2 > max_chars:
            chunks.append(current.strip())
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para
    if current.strip():
        chunks.append(current.strip())
    # Guard against very large paragraphs that exceed max_chars
    final: List[str] = []
    for chunk in chunks:
        while len(chunk) > max_chars:
            # Find last space within limit
            split_at = chunk.rfind(" ", 0, max_chars)
            if split_at == -1:
                split_at = max_chars
            final.append(chunk[:split_at].strip())
            chunk = chunk[split_at:].strip()
        if chunk:
            final.append(chunk)
    return final


def _extract_pdf_pages(data: bytes) -> List[str]:
    """One string per page, so extraction quality can be judged page by page."""
    try:
        import pdfplumber  # type: ignore[import-untyped]

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return [p.extract_text() or "" for p in pdf.pages]
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader  # type: ignore[import-untyped]

        return [p.extract_text() or "" for p in PdfReader(io.BytesIO(data)).pages]
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "PDF parsing requires pdfplumber or PyPDF2. "
                "Install one with: pip install pdfplumber"
            ),
        )


def read_document(data: bytes, ext: str) -> tuple[str, str, int]:
    """A document's text, re-reading any pages the extractor mangled.

    Returns the text, a note describing anything the caller should say out
    loud, and how many pages had to be re-read. A silently partial or garbled
    read is the failure this guards: the extractor returns plenty of
    characters either way, so nothing downstream can tell a clean page from
    one with table rows spliced into the prose.
    """
    if ext in (".txt", ".md", ".csv"):
        try:
            return data.decode("utf-8"), "", 0
        except UnicodeDecodeError:
            return data.decode("latin-1"), "", 0
    if ext == ".docx":
        return _extract_text_from_docx(data), "", 0
    if ext != ".pdf":
        return "", "", 0

    pages = _extract_pdf_pages(data)
    judged = document_read.assess(pages)
    if not any(page.needs_eyes for page in judged):
        return document_read.assemble(judged), "", 0

    replacements, note = document_read.read_pages_with_vision(data, judged)
    return document_read.assemble(judged, replacements), note, len(replacements)


def _extract_text_from_pdf(data: bytes) -> str:
    """Extract text from a PDF using pdfplumber or PyPDF2."""
    # Try pdfplumber first
    try:
        import pdfplumber  # type: ignore[import-untyped]

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n\n".join(pages)
    except ImportError:
        pass

    # Fall back to PyPDF2
    try:
        from PyPDF2 import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n\n".join(pages)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "PDF parsing requires pdfplumber or PyPDF2. "
                "Install one with: pip install pdfplumber"
            ),
        )


def _extract_text_from_docx(data: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            ),
        )


def _get_store() -> KnowledgeStore:
    """Return a KnowledgeStore pointing at the default knowledge DB."""
    db_path = DEFAULT_CONFIG_DIR / "knowledge.db"
    return KnowledgeStore(db_path=db_path)


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------


class PasteRequest(BaseModel):
    title: str = ""
    content: str


class IngestResponse(BaseModel):
    chunks_added: int
    source: str = "upload"


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@router.post("/ingest", response_model=IngestResponse)
async def ingest_paste(body: PasteRequest) -> IngestResponse:
    """Ingest pasted text into the knowledge store."""
    text = body.content.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Content is empty")

    store = _get_store()
    doc_id = str(uuid.uuid4())
    chunks = _chunk_text(text)

    for idx, chunk in enumerate(chunks):
        store.store(
            chunk,
            source="upload",
            doc_type="paste",
            doc_id=doc_id,
            title=body.title or "Pasted text",
            chunk_index=idx,
        )

    logger.info("Ingested %d chunks from pasted text (doc_id=%s)", len(chunks), doc_id)
    return IngestResponse(chunks_added=len(chunks))


class AttachedDocument(BaseModel):
    """One document read for a single conversation, never indexed."""

    filename: str
    text: str
    chars: int
    pages_reread: int = 0
    note: str = ""


@router.post("/attach", response_model=AttachedDocument)
async def attach_document(file: UploadFile = File(...)) -> AttachedDocument:
    """Read one document for the current conversation without storing it.

    The ephemeral half of an attachment: nothing reaches ``knowledge.db``, so a
    one-off file does not silently join the searchable corpus. Indexing is the
    other, explicit choice and goes through ``/ingest/files``.
    """
    filename = file.filename or "untitled"
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in _ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(_ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Allowed: {allowed}",
        )

    data = await file.read()
    if len(data) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(
            status_code=413,
            detail=(
                f"{filename} is larger than the "
                f"{MAX_ATTACHMENT_BYTES // (1024 * 1024)} MB limit."
            ),
        )

    # Off the event loop, always. Reading a mangled paper takes minutes, and
    # run inline it froze the entire server: measured 12 of 12 `/health`
    # probes timing out during one attach, which took the wake-word socket,
    # Flux and TTS down with it and left the user refreshing the page. The
    # identical mistake was diagnosed in `digest_routes` earlier the same day.
    text, note, reread = await asyncio.to_thread(read_document, data, ext)
    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail=(
                f"Nothing readable came out of {filename}. It may be empty, or "
                "an image-only document that could not be read."
            ),
        )

    return AttachedDocument(
        filename=filename,
        text=text,
        chars=len(text),
        pages_reread=reread,
        note=note.strip(),
    )


@router.post("/ingest/files", response_model=IngestResponse)
async def ingest_files(
    files: List[UploadFile] = File(...),
    title: Optional[str] = Form(None),
) -> IngestResponse:
    """Ingest uploaded files into the knowledge store."""
    store = _get_store()
    total_chunks = 0

    for upload in files:
        filename = upload.filename or "untitled"
        ext = ""
        if "." in filename:
            ext = "." + filename.rsplit(".", 1)[-1].lower()

        if ext not in _ALLOWED_EXTENSIONS:
            allowed = ", ".join(sorted(_ALLOWED_EXTENSIONS))
            raise HTTPException(
                status_code=400,
                detail=(f"Unsupported file type: {ext}. Allowed: {allowed}"),
            )

        data = await upload.read()

        # Same reason as `/attach`: this can spend minutes in the vision path.
        text, _note, _reread = await asyncio.to_thread(read_document, data, ext)

        text = text.strip()
        if not text:
            continue

        doc_id = str(uuid.uuid4())
        doc_title = title or filename
        chunks = _chunk_text(text)

        for idx, chunk in enumerate(chunks):
            store.store(
                chunk,
                source="upload",
                doc_type=ext.lstrip("."),
                doc_id=doc_id,
                title=doc_title,
                chunk_index=idx,
            )

        total_chunks += len(chunks)
        logger.info(
            "Ingested %d chunks from file %s (doc_id=%s)",
            len(chunks),
            filename,
            doc_id,
        )

    return IngestResponse(chunks_added=total_chunks)
